import os
import zipfile
from docx import Document

def create_final_doc(content, header, filename_base, token_id):

    filename = f"{filename_base}.docx"

    # 1. Create base DOCX
    doc = Document()
    doc.add_heading(header, level=1)
    doc.add_paragraph(content)
    doc.save(filename)

    # 2. Unzip
    unzip_dir = "temp_unzip"
    os.makedirs(unzip_dir, exist_ok=True)

    with zipfile.ZipFile(filename, 'r') as zip_ref:
        zip_ref.extractall(unzip_dir)

    # 3. Inject canary XML
    rels_path = os.path.join(unzip_dir, "word/_rels/document.xml.rels")
    with open(rels_path, "a", encoding="utf-8") as f:
        f.write(f"""
<Relationship Id="{token_id}"
Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
Target="http://192.120.1.124/tokens/{token_id}.png"
TargetMode="External"/>
""")

    # 4. Rezip
    final_file = f"{filename_base}.docx"
    with zipfile.ZipFile(final_file, 'w') as zipf:
        for root, dirs, files in os.walk(unzip_dir):
            for file in files:
                full_path = os.path.join(root, file)
                arcname = os.path.relpath(full_path, unzip_dir)
                zipf.write(full_path, arcname)

    return final_file
